from flask import render_template, request, send_file, redirect, url_for, session, Response
from docxtpl import DocxTemplate
from docx import Document
import datetime
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx2pdf import convert
import os
import mammoth
from bson import ObjectId
import logging
import pythoncom
import bcrypt
from app import app, db, fs
# Constants
COVER_LETTER_TEMPLATE = os.path.join(app.root_path,'Cover_letterr.docx')
FINAL_FILE_DOCX_FILENAME = "Final_Cover_letter_with_table_{}.docx"
FINAL_FILE_PDF_FILENAME = "Final_Cover_letter_with_table_{}.pdf"
PDFS_DIRECTORY = os.path.join(app.root_path, 'app', 'pdfs')

def generate_cover_letter(context):
    try:
        today_date = datetime.datetime.today().strftime('%B %d, %Y')
        context['today_date'] = today_date

        doc = DocxTemplate(COVER_LETTER_TEMPLATE)
        doc.render(context)
        
        temp_filename = "Temp_Cover_letter.docx"
        doc.save(temp_filename)

        return temp_filename
    except Exception as e:
        logging.error("Error in generate_cover_letter: %s", e)
        raise

def create_and_insert_table(doc, target_index, records):
    try:
        num_cols = 5
        table = doc.add_table(rows=len(records) + 1, cols=num_cols)
        headers = ["S.no", "Description", "Rate", "Quantity", "Amount"]
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header_text
            shading_color = "808080"
            cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in table.rows:
            row.cells[0].width = Inches(0.3)
            row.cells[1].width = Inches(3.0)
            row.cells[2].width = Inches(0.5)
            row.cells[3].width = Inches(0.5)
            row.cells[4].width = Inches(0.5)

        for i, record in enumerate(records, start=1):
            for j, header_text in enumerate(headers):
                cell = table.cell(i, j)
                if header_text == "Amount":
                    rate = float(record[2])
                    quantity = float(record[3])
                    amount = rate * quantity
                    cell.text = str(amount)
                else:
                    cell.text = str(record[j])
                
                shading_color = "D3D3D3" if i % 2 == 0 else "FFFFFF"
                cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.style = 'Table Grid'
        paragraph = doc.paragraphs[target_index]
        paragraph.insert_paragraph_before()._p.addnext(table._tbl)
    except Exception as e:
        logging.error("Error in create_and_insert_table: %s", e)
        raise

def convert_docx_to_html(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            return html
    except Exception as e:
        logging.error("Error in convert_docx_to_html: %s", e)
        raise

@app.route('/', methods=['GET', 'POST'])
def index():
    try:
        if 'username' not in session:
            return redirect(url_for('login'))
            
        if request.method == 'POST':
            context = {
                'offer_name': request.form['offer_name'],
                'offer_no': request.form['offer_no'],
                'company_name': request.form['company_name'],
                'city_name': request.form['city_name'],
                'state_name': request.form['state_name'],
                'manager_name': request.form['manager_name'],
                'enquiry_sub': request.form['enquiry_sub'],
                'delivery_dates': request.form['delivery_dates'],
                'your_name': request.form['your_name'],
                'contact_no': request.form['contact_no']
            }
            records_count = int(request.form['records_count'])
            records = []
            for i in range(records_count):
                sn = request.form[f'sn_{i}']
                description = request.form[f'description_{i}']
                rate = float(request.form[f'rate_{i}'])
                quantity = float(request.form[f'quantity_{i}'])
                records.append((sn, description, rate, quantity))

            cover_letter_file = generate_cover_letter(context)
            doc = Document(cover_letter_file)

            target_text = "Annexure II-Commercial Terms and Conditions."
            target_index = None
            for i, paragraph in enumerate(doc.paragraphs):
                if target_text in paragraph.text:
                    target_index = i
                    break

            if target_index is not None:
                create_and_insert_table(doc, target_index, records)
                unique_suffix = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
                final_docx_file = FINAL_FILE_DOCX_FILENAME.format(unique_suffix)
                doc.save(final_docx_file)

                # Ensure the DOCX file is created before proceeding
                if not os.path.exists(final_docx_file):
                    logging.error("Error: Final DOCX file not created.")
                    return "Error: Final DOCX file not created."

                # Initialize COM library for converting DOCX to PDF
                pythoncom.CoInitialize()
                try:
                    final_pdf_file = os.path.join(PDFS_DIRECTORY, FINAL_FILE_PDF_FILENAME.format(unique_suffix))
                    convert(final_docx_file, final_pdf_file)
                finally:
                    pythoncom.CoUninitialize()

                # Ensure the PDF file is created before proceeding
                if not os.path.exists(final_pdf_file):
                    logging.error("Error: Final PDF file not created.")
                    return "Error: Final PDF file not created."

                # Convert DOCX to HTML for preview
                html_content = convert_docx_to_html(final_docx_file)

                # Save the PDF file to the database
                with open(final_pdf_file, 'rb') as f:
                    file_id = fs.put(f, filename=FINAL_FILE_PDF_FILENAME.format(unique_suffix), username=session['username'])

                return render_template('preview.html', html_content=html_content, file_id=str(file_id))
            else:
                return "Error: Target paragraph not found in the document."
    except Exception as e:
        logging.error("Error in index route: %s", e)
        return str(e)

    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password'].encode('utf-8')

        user = db.users.find_one({"username": username})

        if user and bcrypt.checkpw(password, user['password']):
            session['username'] = username
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error="Invalid username or password.")

    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

@app.route('/download/<file_id>', methods=['GET'])
def download(file_id):
    try:
        file_obj = fs.get(ObjectId(file_id))
        return Response(
            file_obj.read(),
            mimetype='application/pdf',
            headers={"Content-Disposition": f"attachment;filename={file_obj.filename}"}
        )
    except Exception as e:
        logging.error("Error in download route: %s", e)
        return str(e)
    
@app.route('/serve_pdf/<file_id>')
def serve_pdf(file_id):
    try:
        file_obj = fs.get(ObjectId(file_id))

        return Response(
            file_obj.read(),
            mimetype='application/pdf',
            headers={"Content-Disposition": f"inline;filename={file_obj.filename}"}
        )
    except Exception as e:
        logging.error("Error in serve_pdf route: %s", e)
        return str(e)

@app.route('/list_pdfs')
def list_pdfs():
    try:
        if 'username' not in session:
            return redirect(url_for('login'))

        user = db.users.find_one({"username": session['username']})
        is_admin = user.get('is_admin', False)
        
        if is_admin:
            pdf_files = fs.find({"filename": {"$regex": r'\.pdf$'}})
        else:
            pdf_files = fs.find({"filename": {"$regex": r'\.pdf$'}, "username": session['username']})
        
        # Fetching usernames for admin
        if is_admin:
            pdf_files_with_usernames = []
            for pdf in pdf_files:
                username = db.users.find_one({"username": pdf.get('username', '')})
                pdf['username'] = username.get('username', 'Unknown')
                pdf_files_with_usernames.append(pdf)
            return render_template('list_pdfs.html', pdf_files=pdf_files_with_usernames, is_admin=is_admin)
        else:
            return render_template('list_pdfs.html', pdf_files=pdf_files, is_admin=is_admin)
    except Exception as e:
        logging.error("Error in list_pdfs route: %s", e)
        return str(e)

@app.route('/view/<file_id>')
def view_pdf(file_id):
    try:
        file_obj = fs.get(ObjectId(file_id))
        user = db.users.find_one({"username": session['username']})
        is_admin = user.get('is_admin', False)
        return render_template('view_pdf.html', file_id=file_id, filename=file_obj.filename, username=file_obj.username if is_admin else None)
    except Exception as e:
        logging.error("Error in view_pdf route: %s", e)
        return str(e)